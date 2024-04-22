import locale
import pprint

from django.shortcuts import render
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

from .forms import *
from .utils import *
import docx
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE


def akt_constructor(request):
    if request.method == 'POST':
        dolgnost = request.POST.get('dolgnost').replace('\r', '')
        bugor = request.POST.get('bugor').replace('\r', '')
        city = request.POST.get('city').strip()
        commission = dict()
        for i in range(1, 5):
            fio = request.POST.get(f'fio{i}')
            dt = request.POST.get(f'dt{i}')
            if fio:
                commission = {**commission, **{fio: dt}}

        target = request.POST.get('target').replace('\r', '').replace('\n', '')
        lesn = request.POST.get('lesn').strip()
        uch_lesn = request.POST.get('uch_lesn').strip()
        uch = request.POST.get('uch').strip()
        region = request.POST.get('region').strip()
        district = request.POST.get('district').strip()
        lh_osob = request.POST.get('lh_osob').strip()
        ozz = request.POST.get('ozz').strip()
        predl = request.POST.get('predl').strip()
        baza_glr = create_database(request.FILES["data_glr"])
        baza_obsl = create_database(request.FILES["data_obsl"])

        database_validator(baza_glr)
        database_validator(baza_obsl)

        list_kv_obsl = get_list_num_kv(baza_obsl)
        list_kv_glr = get_list_num_kv(baza_glr)

        if list_kv_glr == list_kv_obsl:
            list_num_kv = get_list_num_kv(baza_obsl)
            doc = docx.Document()

            section = doc.sections[0]
            section.left_margin = Mm(25)
            section.right_margin = Mm(15)
            section.top_margin = Mm(15)
            section.bottom_margin = Mm(15)

            st = doc.styles['Normal']
            st.font.name = 'Times New Roman'
            st.font.size = Pt(12)
            st.paragraph_format.space_after = Mm(0)

            for num_kv in list_num_kv:
                list_vd_glr = filter_kv(baza_glr, num_kv)
                list_vd_obsl = filter_kv(baza_obsl, num_kv)

                str_vd = get_str_num_vd(list_vd_obsl)
                square_kv = round(sum(vd[1][2] for vd in list_vd_obsl), 1)
                square_lesn = round(sum((vd[1][2] if vd[1][3] <= 74 else 0.0) for vd in list_vd_obsl), 1)
                square_nles = round(sum((vd[1][2] if vd[1][3] >= 82 else 0.0) for vd in list_vd_obsl), 1)
                square_pok = round(sum((vd[1][2] if vd[1][3] <= 51 else 0.0) for vd in list_vd_obsl), 1)
                square_npok = round(sum((vd[1][2] if 53 <= vd[1][3] <= 74 else 0.0) for vd in list_vd_obsl), 1)
                square_nlk = round(sum((vd[1][2] if 31 <= vd[1][3] <= 37 else 0.0) for vd in list_vd_obsl), 1)
                square_pash = round(sum((vd[1][2] if vd[1][3] == 82 else 0.0) for vd in list_vd_obsl), 1)
                square_sen = round(sum((vd[1][2] if vd[1][3] == 83 else 0.0) for vd in list_vd_obsl), 1)
                square_past = round(sum((vd[1][2] if vd[1][3] == 84 else 0.0) for vd in list_vd_obsl), 1)
                square_vod = round(sum((vd[1][2] if 101 <= vd[1][3] <= 109 else 0.0) for vd in list_vd_obsl), 1)
                square_dp = round(sum((vd[1][2] if 121 <= vd[1][3] <= 141 else 0.0) for vd in list_vd_obsl), 1)
                square_bol = round(sum((vd[1][2] if vd[1][3] == 217 else 0.0) for vd in list_vd_obsl), 1)
                square_proch = round(sum((vd[1][2] if vd[1][3] == 239 else 0.0) for vd in list_vd_obsl), 1)

                doc.add_paragraph(f'Утверждаю\n{dolgnost}{bugor}\n\n{"_" * 16}\nм.п.').alignment = WD_ALIGN_PARAGRAPH.RIGHT
                doc.add_paragraph(f'Акт № {"_" * 7}').alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph('несоответствия данных государственного '
                                  'лесного реестра натурному обследованию\n').alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph(f'{city}{" " * (90 - len(city))}{"_" * 15}\n').alignment = WD_ALIGN_PARAGRAPH.CENTER
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                chlens = ', '.join([' '.join(list(chlen)) for chlen in commission.items()])
                p.add_run(chlens).underline = WD_UNDERLINE.SINGLE
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.add_run(f'провели натурное обследование лесного участка, в целях: ')
                p.add_run(f'{target}').underline = WD_UNDERLINE.SINGLE
                doc.add_paragraph(f'При обследовании уточнены данные лесного реестра и установлено:')
                p = doc.add_paragraph()

                # 1.
                p.add_run(f'1. Участок расположен в лесах ')
                p.add_run(lesn).underline = WD_UNDERLINE.SINGLE
                p.add_run(f', {uch_lesn}, {uch}, в квартале №{num_kv} выделах {str_vd}')
                p = doc.add_paragraph()
                p.add_run(f'Субъект Российской Федерации ')
                p.add_run(f'{region}').underline = WD_UNDERLINE.SINGLE
                p = doc.add_paragraph()
                p.add_run(f'Муниципальный район ')
                p.add_run(f'{district}').underline = WD_UNDERLINE.SINGLE
                p = doc.add_paragraph()

                # 2.
                p.add_run(f'2. Подразделение лесов по целевому назначению ')
                p.add_run(f'эксплуатационные леса').underline = WD_UNDERLINE.SINGLE
                p = doc.add_paragraph()

                # 3.
                p.add_run(f'3. Категория защитных лесов ')
                p.add_run(f' - ').underline = WD_UNDERLINE.SINGLE
                p = doc.add_paragraph()

                # 4.
                p.add_run(f'4. Общая площадь участка ')
                p.add_run(f'{str(square_kv).replace(".", ",")}').underline = WD_UNDERLINE.SINGLE
                p.add_run(f' га.,')
                doc.add_paragraph('в том числе:')
                p = doc.add_paragraph()
                p.add_run(f'лесных земель ')
                p.add_run(f'{str(square_lesn).replace(".", ",")}').underline = WD_UNDERLINE.SINGLE
                p.add_run(f' га.')
                p = doc.add_paragraph()
                p.add_run(f'из них: покрытых лесной растительностью ')
                p.add_run(f'{str(square_pok).replace(".", ",")}').underline = WD_UNDERLINE.SINGLE
                p.add_run(f' га.')
                p = doc.add_paragraph()
                p.add_run(f'{" " * 13}не покрытых лесом ')
                p.add_run(f'{str(square_npok).replace(".", ",")}').underline = WD_UNDERLINE.SINGLE
                p.add_run(f' га.')
                p = doc.add_paragraph()
                p.add_run(f'{" " * 13}в том числе несомкнувшихся лесных культур ')
                p.add_run(f'{str(square_nlk).replace(".", ",")}').underline = WD_UNDERLINE.SINGLE
                p.add_run(f' га.')
                p = doc.add_paragraph()
                p.add_run(f'\nнелесных земель ')
                p.add_run(f'{str(square_nles).replace(".", ",")}').underline = WD_UNDERLINE.SINGLE
                p.add_run(f' га.')
                p = doc.add_paragraph()
                p.add_run(f'из них: пашен ')
                p.add_run(f'{str(square_pash).replace(".", ",")}').underline = WD_UNDERLINE.SINGLE
                p.add_run(f' га.')
                p = doc.add_paragraph()
                p.add_run(f'{" " * 13}сенокосов ')
                p.add_run(f'{str(square_sen).replace(".", ",")}').underline = WD_UNDERLINE.SINGLE
                p.add_run(f' га.')
                p = doc.add_paragraph()
                p.add_run(f'{" " * 13}пастбищ ')
                p.add_run(f'{str(square_past).replace(".", ",")}').underline = WD_UNDERLINE.SINGLE
                p.add_run(f' га.')
                p = doc.add_paragraph()
                p.add_run(f'{" " * 13}вод ')
                p.add_run(f'{str(square_vod).replace(".", ",")}').underline = WD_UNDERLINE.SINGLE
                p.add_run(f' га.')
                p = doc.add_paragraph()
                p.add_run(f'{" " * 13}дорог, просек ')
                p.add_run(f'{str(square_dp).replace(".", ",")}').underline = WD_UNDERLINE.SINGLE
                p.add_run(f' га.')
                p = doc.add_paragraph()
                p.add_run(f'{" " * 13}болот ')
                p.add_run(f'{str(square_bol).replace(".", ",")}').underline = WD_UNDERLINE.SINGLE
                p.add_run(f' га.')
                str(square_bol).replace(',', '.')
                p = doc.add_paragraph()
                p.add_run(f'{" " * 13}прочих земель ')
                p.add_run(f'{str(square_proch).replace(".", ",")}').underline = WD_UNDERLINE.SINGLE
                p.add_run(f' га.')

                # 5.
                doc.add_paragraph('\n5. Таксационное описание по материалам государственного лесного реестра:\n')
                create_table(doc, list_vd_glr, lesn, uch_lesn, uch)

                # 6.
                doc.add_paragraph('\n6. Таксационное описание по результатам обследования:\n')
                create_table(doc, list_vd_obsl, lesn, uch_lesn, uch)

                # 7.
                p = doc.add_paragraph()
                p.add_run(f'\n7. Участок ')
                p.add_run(ozz).underline = WD_UNDERLINE.SINGLE
                p.add_run(f' особо защитное значение')

                # 8.
                p = doc.add_paragraph()
                p.add_run(f'\n8. Лесохозяйственные особенности участка: ')
                p.add_run(lh_osob).underline = WD_UNDERLINE.SINGLE

                # 9.
                p = doc.add_paragraph()
                p.add_run(f'\n9. При составлении акта сделаны следующие замечания и предложения: ')
                p.add_run(predl).underline = WD_UNDERLINE.SINGLE

                p = doc.add_paragraph()
                p.add_run(f'\nЛица, проводившие обследование:')

                for chlen in commission:
                    p = doc.add_paragraph()
                    p.add_run(f'\n\n{chlen}{" " * (135 - len(chlen))}.').underline = WD_UNDERLINE.SINGLE
                    doc.add_paragraph(f'(Ф.И.О., подпись)').alignment = WD_ALIGN_PARAGRAPH.CENTER

                doc.add_paragraph(f'\n\nНеотъемлемой частью настоящего акта является чертеж лесного участка')

                doc.add_page_break()

            doc.save('Акт несоответствия.docx')

        form = AktForm(request.POST, request.FILES)
    else:
        form = AktForm()
    context = {
        'form': form,
    }
    return render(request, "Akt_NV/akt_constructor.html", context=context)
