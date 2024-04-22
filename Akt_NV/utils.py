import math
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Mm, Cm


def create_database(file):
    """
    Построить базу из текстового файла в кодировке cp866
    :param file:
    :return:
    """
    st = file.read()
    st = st.decode('cp866')  # Раскодировать строку
    st = st.replace('\r\n', '')  # Удалить лишние символы
    vd_list_stroka = st.split('?')  # Список выделов (строки)
    vd_list_stroka.pop(0)  # Удалить мусор
    vd_list_stroka.pop()
    baza = []
    # База список выделов - список макетов - список позиций
    for vd_stroka in vd_list_stroka:
        maket_list_stroka = vd_stroka.split(':')
        vd_lst = []
        for maket_stroka in maket_list_stroka:
            vd_lst.append(maket_stroka.split('.'))
        baza.append(vd_lst)

    # Дописать нулевой макет
    maket_0 = ''
    for vd in baza:
        if vd[0][0] == '0':
            maket_0 = vd[0]
        else:
            vd.insert(0, maket_0)
    return baza


def database_validator(baza):
    for vd in baza:
        for maket in vd:
            for i in maket:
                maket[0] = int(maket[0])
        vd[0][1] = int(vd[0][1])
        vd[0][2] = int(vd[0][2])
        vd[1][2] = float(vd[1][2].replace(',', '.'))
        vd[1][3] = int(vd[1][3])


def get_line_numbers_concat(line_nums):
    """
    Список номеров в короткую строку
    :param line_nums:
    :return:
    """
    seq = []
    final = []
    last = 0

    for index, val in enumerate(line_nums):

        if last + 1 == val or index == 0:
            seq.append(val)
            last = val
        else:
            if len(seq) > 1:
                final.append(str(seq[0]) + '-' + str(seq[len(seq) - 1]))
            else:
                final.append(str(seq[0]))
            seq = []
            seq.append(val)
            last = val

        if index == len(line_nums) - 1:
            if len(seq) > 1:
                final.append(str(seq[0]) + '-' + str(seq[len(seq) - 1]))
            else:
                final.append(str(seq[0]))

    final_str = ', '.join(map(str, final))
    return final_str


def filter_kv(baza, num_kv):
    sample = []
    for vd in baza:
        if vd[0][2] == num_kv:
            sample.append(vd)
    return sample


def get_str_num_vd(lst):
    """
    Сформировать список выделов в кварале
    :param lst:
    :return:
    """
    vd_list = []
    for vd in lst:
        vd_list.append(int(vd[1][1]))
    vd_list.sort()
    return get_line_numbers_concat(vd_list)


def get_list_num_kv(baza):
    """
    Сформировать список кварталов
    :param baza:
    :return:
    """
    s = set()
    for vd in baza:
        s.add(int(vd[0][2]))
    return list(s)


def get_sostav(vd):
    sostav = ''
    for maket in vd:
        if maket[0] == 10:
            if maket[1]:
                if sostav:
                    sostav += ' '
                if int(maket[1]) == 9:
                    sostav += 'Ед'
            if maket[2]:
                sostav += maket[2]
            else:
                sostav += ','
            sostav += maket[3]
    return sostav.replace(',', '+', 1)


def get_zapas_kv(list_vd):
    zapas = 0
    for vd in list_vd:
        for maket in vd:
            if maket[0] == 10:
                zapas += int(maket[11])
                break
    return zapas


def get_bonitet(vd):
    for maket in vd:
        if maket[0] == 3:
            return maket[2]


def get_polnota(vd):
    for maket in vd:
        if maket[0] == 10:
            return maket[9].replace(".", ",")


def get_zapas(vd):
    for maket in vd:
        if maket[0] == 10:
            return maket[11]


def get_vozrast(vd):
    for maket in vd:
        if maket[0] == 10:
            kl_10 = ['Б', 'ОС', 'ОЛС', 'ОЛЧ', ]
            kl_20 = ['С', 'П', 'Е', 'Л', ]
            kl_40 = ['К', ]
            if maket[3] in kl_10:
                return f'{math.ceil(int(maket[4]) / 10)}/{maket[4]}'
            if maket[3] in kl_20:
                return f'{math.ceil(int(maket[4]) / 20)}/{maket[4]}'
            if maket[3] in kl_40:
                return f'{math.ceil(int(maket[4]) / 40)}/{maket[4]}'
            return f'{maket[4]}'


kz_handbook = {
    3: 'Насаждение естественного происхождения',
    4: 'Насаждение из подроста',
    5: 'Насаждение, расстроенное условно-сплошными рубками',
    7: 'Насаждение в стадии реконструкции',
    8: 'Насаждени, имеющее в составе примесь пород искусственного происхождения',
    9: 'Насаждение с лесными культурами под пологом',
    10: 'Лесные культуры',
    11: 'Культуры, созданные путём реконструкции',
    12: 'Культуры видовые',
    13: 'Культуры декоративные',
    14: 'Культуры ландшафтные',
    15: 'Культуры - полезащитные лесные полосы',
    16: 'Культуры с культурами под пологом',
    17: 'Культуры в стадии реконструкции',
    31: 'Несомкнувшиеся культуры',
    32: 'Несомкнувшиеся культуры, созданные путём реконструкции',
    34: 'Несомкнувшиеся культуры видовые',
    35: 'Несомкнувшиеся культуры декоративные',
    36: 'Несомкнувшиеся культуры ландшафтные',
    37: 'Несомкнувшиеся культуры - полезащитные лесные полосы',
    40: 'Питомник',
    41: 'Плантация',
    42: 'Школа древесная',
    43: 'Дендросад',
    44: 'Теплица',
    51: 'Редина естественная',
    53: 'Гарь',
    54: 'Погибшее насаждение',
    55: 'Ветровальник',
    57: 'Лесосека года лесоустройства',
    58: 'Раскорчёвванная площадь',
    59: 'Вырубка',
    60: 'Прогалина',
    61: 'Пустырь',
    62: 'Земли рекультивированные',
    82: 'Пашня',
    83: 'Сенокос',
    84: 'Пастбище, выгон',
    85: 'Сырт',
    86: 'Луг субальпийский, альпийский',
    87: 'Скотопрогон',
    90: 'Поле кормовое',
    101: 'Озеро',
    102: 'Река',
    103: 'Ручей',
    104: 'Пруд',
    105: 'Водохранилище',
    106: 'Канал',
    107: 'Канава мелиоративная, арык',
    108: 'Сеть коллекторная',
    109: 'Старица',
    110: 'Сад (не дендрологический)',
    111: 'Коллективный сад',
    112: 'Тутовник',
    115: 'Ягодник культурный',
    116: 'Хмельник',
    121: 'Дорога железная',
    122: 'Дорога железная узкой колеи',
    123: 'Дорога автомобильная с искусственным покрытием',
    125: 'Дорога автомобильная грунтовая',
    126: 'Зимник',
    127: 'Тропа',
    128: 'Дорожка пешеходная',
    129: 'Просека квартальная',
    130: 'Коридор технологический',
    131: 'Визир',
    132: 'Граница окружная',
    133: 'Канал (не мелиоративный)',
    140: 'Разрыв противопожарный',
    141: 'Прогон',
    162: 'Кордон',
    163: 'Усадьба частная',
    164: 'Склад лесной',
    165: 'Пасека пчелиная',
    166: 'Станция метеорологическая',
    167: 'Площадка вертолётная',
    168: 'Посёлок лесной',
    169: 'Зимовье, охотничий домик',
    170: 'Кемпинг',
    171: 'Дом отдыха',
    172: 'Лагерь отдыха (пионерский)',
    174: 'Сооружение парковое',
    175: 'Площадка детская',
    176: 'Стадион',
    177: 'Площадка спортивная или игровая',
    178: 'Площадка с памятником',
    182: 'Пляж',
    183: 'Стоянка автотранспорта',
    184: 'Санаторий',
    185: 'Кладбище',
    187: 'Скотомогильник',
    188: 'Ремиз',
    189: 'Поляна для отдыха',
    190: 'Ландшафтная поляна',
    211: 'Овраг',
    212: 'Балка',
    213: 'Склон крутой',
    214: 'Обнажение скальное',
    215: 'Россыпь каменистая, галечник',
    216: 'Пески',
    217: 'Болото',
    218: 'Марь',
    219: 'Плавни (не камышовые)',
    220: 'Солодь',
    221: 'Солонец',
    222: 'Гольцы',
    223: 'Тундра (не горная)',
    224: 'Горная тундра',
    225: 'Снежник',
    226: 'Ледник',
    227: 'Отвал',
    228: 'Нефтекачалка',
    229: 'Плавни камышовые',
    230: 'Осыпь',
    231: 'Оползень',
    232: 'Образование карстовое',
    233: 'Наледь',
    234: 'Спецплощадка',
    239: 'Прочие земли',
    240: 'Карьер действующий',
    241: 'Торфоразработки',
    242: 'Отвал пород, золы, шлака',
    248: 'Линия (трасса) электропередач',
    249: 'Линия связи',
    250: 'Газопровод',
    251: 'Нефтепровод',
    252: 'Водопровод',
    253: 'Прочие трассы',
    254: 'Мусоропровод',
    255: 'Свалка мусора и производственных отходов',
    256: 'Трасса мелиоративная',
}


def create_table(doc, list_vd, lesn, uch_lesn, uch):
    # Создать таблицу
    table = doc.add_table(len(list_vd) + 3, 9)
    table.style = 'Table Grid'
    table.autofit = False

    # Внести данные в таблицу
    rows = table.rows
    table_head = ('Наименование лесничества', '№ квартала', '№ выдела', 'Площадь, га', 'Состав насаждения',
                  'Класс возраста/ возраст, лет', 'Бонитет', 'Полнота', 'Общий запас древесины куб. м',
                  )

    for ic in range(9):
        # Заполнить заголовок
        cell = rows[0].cells[ic]
        cell.paragraphs[0].add_run(table_head[ic])

        # Заполнить вторую строку
        cell = rows[1].cells[ic]
        cell.paragraphs[0].add_run(str(ic + 1))

    # Заполнить две левые колонки
    cell = rows[2].cells[0]
    cell.paragraphs[0].add_run(f'{lesn} {uch_lesn} {uch}')

    cell = rows[2].cells[1]
    cell.paragraphs[0].add_run(str(list_vd[0][0][2]))

    # Внести данные из базы по кварталу
    for ir, row in enumerate(rows[2:-1]):
        for ic, cell in enumerate(row.cells[2:]):
            p = cell.paragraphs[0]
            if ic == 0:
                p.add_run(list_vd[ir][1][1])
            elif ic == 1:
                p.add_run(str(list_vd[ir][1][2]).replace(".", ","))
            elif ic == 2:
                sostav = get_sostav(list_vd[ir])
                if sostav:
                    p.add_run(sostav)
                else:
                    p.add_run(kz_handbook.get(list_vd[ir][1][3]))
            elif ic == 3:
                p.add_run(get_vozrast(list_vd[ir]))
            elif ic == 4:
                p.add_run(get_bonitet(list_vd[ir]))
            elif ic == 5:
                p.add_run(get_polnota(list_vd[ir]))
            elif ic == 6:
                p.add_run(get_zapas(list_vd[ir]))

    row = rows[-1]
    # Итог
    cell = row.cells[2]
    cell.paragraphs[0].add_run('Итого')
    # Площадь квартала
    cell = row.cells[3]
    cell.paragraphs[0].add_run(str(round(sum(vd[1][2] for vd in list_vd), 1)).replace('.', ','))
    # Запас квартала
    cell = row.cells[8]
    cell.paragraphs[0].add_run(str(get_zapas_kv(list_vd)))

    # Форматировать таблицу
    for ir, row in enumerate(rows):
        cells = row.cells
        cells[0].width = Cm(2.70)
        cells[1].width = Cm(1.50)
        cells[2].width = Cm(1.25)
        cells[3].width = Cm(1.80)
        cells[4].width = Cm(4.00)
        cells[5].width = Cm(1.80)
        cells[6].width = Cm(1.50)
        cells[7].width = Cm(1.50)
        cells[8].width = Cm(1.80)
        for ic, cell in enumerate(cells):
            if ir == 0:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.left_indent = Cm(-0.15)
            paragraph.paragraph_format.right_indent = Cm(-0.15)
            for run in paragraph.runs:
                run.font.size = Pt(10)

        # Объединить колонки
        for i in range(2):
            col = table.column_cells(i)
            listOfCellsToMerge = col[2:len(col)]
            listOfCellsToMerge[0].merge(listOfCellsToMerge[-1])
